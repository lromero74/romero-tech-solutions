from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/louis/New/02_Work/JobApplications/CoverLetters/Flex_Staff_Infrastructure_Engineer_Cover_Letter.docx"


def set_run(run, size=10.5, bold=False, color=None):
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Louis Antonio Romero")
set_run(r, 16, True, (31, 78, 121))

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = contact.add_run("Escondido, CA | louis@romerotechsolutions.com | (734) 255-7060 | linkedin.com/in/gpsdirect | github.com/lromero74")
set_run(r, 9.5, False, (70, 70, 70))

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(4)
rule.paragraph_format.space_after = Pt(10)
r = rule.add_run("_" * 98)
set_run(r, 7, False, (180, 190, 200))

for line in [
    "April 29, 2026",
    "",
    "Flex Hiring Team",
    "Re: Staff Infrastructure Engineer",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(line)
    set_run(r, 10.5)

body = [
    "Dear Flex Hiring Team,",
    "I am excited to apply for the Staff Infrastructure Engineer role because it lines up unusually well with the kind of work I do best: setting practical infrastructure direction, improving reliability across production systems, raising the bar for developer experience, and turning operational lessons into durable platform improvements. Flex's mission in rent payment flexibility also sits in a serious operational domain where reliability, security, and user trust matter every day.",
    "Most recently, through Romero Tech Solutions, I have been building and operating a production MSP platform across AWS, React/TypeScript, Express, PostgreSQL, Socket.IO, Stripe, SES, Cognito-style auth flows, RBAC, audit trails, workflow automation, monitoring, and deployment discipline. That work has required the same blend Flex describes: hands-on infrastructure ownership, clear technical decision-making, secure production operations, developer-facing automation, and the judgment to keep AI-assisted engineering useful without giving up human accountability for correctness, security, or reliability.",
    "Earlier roles at Apple, TikTok/ByteDance, Blackhawk Network, TDAmeritrade/Charles Schwab, and ADP gave me deep exposure to build/release reliability, Linux and production operations, cloud infrastructure, incident response, regulated environments, and large-scale support workflows. I am comfortable leading across ambiguity: defining SLOs, improving observability, hardening CI/CD, using Terraform and cloud-native patterns, partnering with product and engineering leaders, and writing the strategy and tradeoff documents needed to align teams.",
    "What stands out about Flex is the combination of Staff-level infrastructure scope and an AI-first engineering culture that still emphasizes strong human ownership. I would bring a bias for practical delivery: identify the highest-leverage reliability and developer-platform investments, build the tooling and standards that make good behavior easier, and mentor engineers toward operational rigor without slowing down product momentum.",
    "Thank you for your consideration. I would welcome the chance to discuss how my platform/SRE background, AWS and production reliability experience, and current AI-assisted engineering work can help Flex scale its infrastructure with confidence.",
    "Sincerely,",
    "Louis Antonio Romero",
]

for i, text in enumerate(body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7 if i in {0, 6} else 0)
    p.paragraph_format.space_after = Pt(8 if text else 0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_run(r, 10.5)

doc.save(OUTPUT)
print(OUTPUT)
